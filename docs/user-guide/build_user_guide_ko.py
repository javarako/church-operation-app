from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

import build_user_guide as base


GUIDE_DIR = Path(__file__).resolve().parent
CONTENT_PATH = GUIDE_DIR / "user-guide-content-ko.md"
OUTPUT_PATH = GUIDE_DIR / "교회운영 메뉴얼.docx"
KOREAN_FONT = "Nanum Gothic"


original_set_run_font = base.set_run_font
original_set_cell_margins = base.set_cell_margins


def set_korean_run_font(run, size=None, color=None, bold=None, italic=None, name=KOREAN_FONT):
    if size == 9.2:
        size = 8.5
    original_set_run_font(run, size, color, bold, italic, name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_compact_cell_margins(cell, top=55, start=100, bottom=55, end=100):
    original_set_cell_margins(cell, top, start, bottom, end)


def add_compact_figure(doc: Document, filename: str, caption: str) -> None:
    image_path = base.SCREENSHOT_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = base.Pt(4)
    p.paragraph_format.space_after = base.Pt(2)
    p.paragraph_format.keep_with_next = True
    auth_screens = {"01-login.png", "02-forgot-password.png", "03-reset-password.png", "12-change-password.png"}
    width = 4.7 if filename in auth_screens else 6.55
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(6)
    set_korean_run_font(p.add_run(caption), size=8.5, color=base.MID_GRAY, italic=True)


def apply_korean_typography(doc: Document) -> None:
    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = KOREAN_FONT
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), KOREAN_FONT)
        r_fonts.set(qn("w:hAnsi"), KOREAN_FONT)
        r_fonts.set(qn("w:eastAsia"), KOREAN_FONT)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = base.Pt(0)
    set_korean_run_font(header.add_run("교회운영  |  사용자 안내서"), size=8.5, color=base.MID_GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_korean_run_font(footer.add_run("Capstone Presbyterian Church   |   "), size=8, color=base.MID_GRAY)
    base.add_field(footer, "PAGE")


def add_cover(doc: Document) -> None:
    if base.LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(base.LOGO_PATH), width=Inches(2.15))
        shape._inline.docPr.set("descr", "Capstone Presbyterian Church 로고")
        p.paragraph_format.space_after = base.Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(8)
    set_korean_run_font(p.add_run("CAPSTONE PRESBYTERIAN CHURCH"), size=11, color=base.TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(8)
    set_korean_run_font(p.add_run("교회운영 메뉴얼"), size=30, color=base.NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(24)
    set_korean_run_font(p.add_run("사용자 안내서"), size=18, color=base.BLUE, bold=True)

    if base.BANNER_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(base.BANNER_PATH), width=Inches(6.6))
        shape._inline.docPr.set("descr", "교회 배너")
        p.paragraph_format.space_after = base.Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(5)
    set_korean_run_font(
        p.add_run("Admin, Treasurer, Pastor, Membership, Viewer 및 Member 역할용"),
        size=10.5,
        color=base.MID_GRAY,
        italic=True,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_korean_run_font(p.add_run("버전 1.0  |  2026년 7월"), size=10, color=base.MID_GRAY, bold=True)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("목차", level=1)
    chapters = [
        "이 안내서 소개", "역할 및 접근 권한", "시작하기", "대시보드", "교인 정보",
        "헌금 관리", "재정 관리", "예산 관리", "기준정보", "보고서", "내 프로필",
        "로그아웃", "문제 해결 및 보안 안내",
    ]
    for chapter in chapters:
        doc.add_paragraph(chapter, style="List Bullet")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = base.Pt(12)
    set_korean_run_font(
        p.add_run("팁: Word의 탐색 창을 사용하면 제목 사이를 빠르게 이동할 수 있습니다."),
        size=9.5,
        color=base.MID_GRAY,
        italic=True,
    )
    doc.add_page_break()


def build() -> None:
    base.CONTENT_PATH = CONTENT_PATH
    base.OUTPUT_PATH = OUTPUT_PATH
    base.set_run_font = set_korean_run_font
    base.set_cell_margins = set_compact_cell_margins
    base.add_figure = add_compact_figure

    doc = Document()
    base.configure_document(doc)
    apply_korean_typography(doc)
    doc.core_properties.title = "교회운영 메뉴얼"
    doc.core_properties.subject = "Church Operations 역할별 한국어 사용자 안내서"
    doc.core_properties.creator = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "교회운영, 교인, 헌금, 재정, 예산, 보고서"
    add_cover(doc)
    add_contents(doc)
    base.parse_content(doc)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
