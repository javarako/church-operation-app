from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
GUIDE_DIR = ROOT / "docs" / "user-guide"
CONTENT_PATH = GUIDE_DIR / "user-guide-content.md"
OUTPUT_PATH = GUIDE_DIR / "Church Operations User Guide.docx"
SCREENSHOT_DIR = GUIDE_DIR / "screenshots"
LOGO_PATH = ROOT / "backend" / "src" / "main" / "resources" / "static" / "branding" / "church_logo.png"
BANNER_PATH = ROOT / "backend" / "src" / "main" / "resources" / "static" / "branding" / "church-banner.png"

NAVY = "123047"
TEAL = "087D8F"
BLUE = "2E5F85"
LIGHT_BLUE = "E8F1F5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "202833"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Aptos") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def create_restarted_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    style_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_num_id = style_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def add_inline_markup(paragraph, text: str) -> None:
    pieces = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, name="Courier New", size=9.5, color=NAVY)
        else:
            paragraph.add_run(piece)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Named override: screenshot_manual_compact_page, allowing wider UI figures.
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    style_tokens = {
        "Heading 1": (18, NAVY, 16, 8),
        "Heading 2": (14, BLUE, 13, 6),
        "Heading 3": (11.5, TEAL, 9, 4),
    }
    for style_name, (size, color, before, after) in style_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    p = header.paragraphs[0]
    p.text = "CHURCH OPERATIONS  |  USER GUIDE"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.runs[0], size=8.5, color=MID_GRAY, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Capstone Presbyterian Church   |   ")
    set_run_font(run, size=8, color=MID_GRAY)
    add_field(p, "PAGE")


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(LOGO_PATH), width=Inches(2.15))
        shape._inline.docPr.set("descr", "Capstone Presbyterian Church logo")
        p.paragraph_format.space_after = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("CAPSTONE PRESBYTERIAN CHURCH")
    set_run_font(run, size=11, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Church Operations")
    set_run_font(run, size=30, color=NAVY, bold=True, name="Aptos Display")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("User Guide")
    set_run_font(run, size=18, color=BLUE, bold=True, name="Aptos Display")

    if BANNER_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(BANNER_PATH), width=Inches(6.6))
        shape._inline.docPr.set("descr", "Church banner")
        p.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("For Admin, Treasurer, Pastor, Membership, Viewer, and Member roles")
    set_run_font(run, size=10.5, color=MID_GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  July 2026")
    set_run_font(run, size=10, color=MID_GRAY, bold=True)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    chapters = [
        "About This Guide", "Roles and Access", "Getting Started", "Dashboard",
        "Member Information", "Offering Management", "Financial Management",
        "Budget Management", "Reference Data", "Reports", "My Profile",
        "Logout", "Troubleshooting and Security Notes",
    ]
    for chapter in chapters:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(chapter)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Tip: In Word, use the Navigation pane to jump between headings.")
    set_run_font(run, size=9.5, color=MID_GRAY, italic=True)
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    width = 7.0 / len(rows[0])

    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(rows[0]):
        cell = header.cells[index]
        cell.width = Inches(width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9.2, color=WHITE, bold=True)

    for row_values in rows[1:]:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, filename: str, caption: str) -> None:
    image_path = SCREENSHOT_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    auth_screens = {"01-login.png", "02-forgot-password.png", "03-reset-password.png", "12-change-password.png"}
    width = 4.8 if filename in auth_screens else 6.75
    shape = run.add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(caption)
    set_run_font(run, size=8.5, color=MID_GRAY, italic=True)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    cell.width = Inches(6.85)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline_markup(p, text)
    for run in p.runs:
        set_run_font(run, size=9.5, color=NAVY, bold=run.bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_content(doc: Document) -> None:
    lines = CONTENT_PATH.read_text(encoding="utf-8").splitlines()
    index = 0
    first_source_title = True
    major_section_index = 0
    active_numbering_id = None
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            active_numbering_id = None
            index += 1
            continue

        if line.startswith("Version "):
            index += 1
            continue

        if line.startswith("[[FIGURE:"):
            match = re.match(r"\[\[FIGURE:(.*?)\|(.*?)\]\]", line)
            if match:
                add_figure(doc, match.group(1), match.group(2))
            index += 1
            continue

        if line.startswith("|" ):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            parsed = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                if all(re.fullmatch(r"[-: ]+", cell) for cell in cells):
                    continue
                parsed.append(cells)
            add_table(doc, parsed)
            continue

        if line.startswith("# "):
            if first_source_title:
                first_source_title = False
            index += 1
            continue

        if line.startswith("## "):
            if major_section_index >= 2:
                doc.add_page_break()
            doc.add_heading(line[3:], level=1)
            major_section_index += 1
            index += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            index += 1
            continue

        if line.startswith("> "):
            add_callout(doc, line[2:])
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            if active_numbering_id is None:
                active_numbering_id = create_restarted_numbering(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_numbering_id)
            add_inline_markup(p, number_match.group(1))
            index += 1
            continue


        active_numbering_id = None

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, line[2:])
            index += 1
            continue

        p = doc.add_paragraph()
        add_inline_markup(p, line)
        index += 1


def build() -> None:
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Church Operations User Guide"
    doc.core_properties.subject = "Role-aware user manual for Church Operations"
    doc.core_properties.creator = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "church operations, members, offerings, finance, budgets, reports"
    add_cover(doc)
    add_contents(doc)
    parse_content(doc)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
